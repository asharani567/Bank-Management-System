from database import create_connection
import tkinter as tk
from tkinter import messagebox
import random
from PIL import Image, ImageDraw, ImageFont,ImageTk
from docx import Document
from datetime import datetime
import os
import mysql.connector
from datetime import datetime,date

class AdminLogin:
	def __init__(self, root, previous_page_callback=None):
		self.root = root
		self.previous_page_callback = previous_page_callback
		self.root.title("Admin Login")
		self.root.attributes('-fullscreen', True)
		self.create_login_ui()

	def create_login_ui(self):
		# Clear existing widgets
		for widget in self.root.winfo_children():
			widget.destroy()

		# Background Image
		bg_image = Image.open("admin_login_bg.jpg")
		bg_image = bg_image.resize((self.root.winfo_screenwidth(), self.root.winfo_screenheight()), Image.LANCZOS)
		self.bg_photo = ImageTk.PhotoImage(bg_image)
		bg_label = tk.Label(self.root, image=self.bg_photo)
		bg_label.place(x=0, y=0, relwidth=1, relheight=1)

		# Login Frame
		login_frame = tk.Frame(self.root, bg="white", padx=30, pady=30, bd=5, relief=tk.RIDGE)
		login_frame.place(relx=0.5, rely=0.5, anchor=tk.CENTER)

		# Title
		tk.Label(login_frame, text="Admin Login", font=("Arial", 24, "bold"), bg="white").pack(pady=20)

		# Username
		tk.Label(login_frame, text="Username:", font=("Arial", 14), bg="white").pack(pady=5)
		self.username_entry = tk.Entry(login_frame, font=("Arial", 14), width=20)
		self.username_entry.pack(pady=5)

		# Password
		tk.Label(login_frame, text="Password:", font=("Arial", 14), bg="white").pack(pady=5)
		self.password_entry = tk.Entry(login_frame, font=("Arial", 14), show="*", width=20)
		self.password_entry.pack(pady=5)

		# Buttons Frame
		button_frame = tk.Frame(login_frame, bg="white")
		button_frame.pack(pady=10)

		tk.Button(button_frame, text="Login", command=self.login, font=("Arial", 14, "bold"),bg="#004080", fg="white", padx=20, pady=10,width=8,height=1).pack(side=tk.LEFT,padx=5,pady=15)
		tk.Button(button_frame, text="Back", command=self.go_back, font=("Arial", 14, "bold"),bg="red", fg="white", padx=20, pady=10, width=8,height=1).pack(side=tk.RIGHT,padx=5,pady=20)


	def login(self):
		username = self.username_entry.get()
		password = self.password_entry.get()
		if username == "admin" and password == "admin":
			messagebox.showinfo("Login Successful", "Welcome, Admin!")
			self.open_admin_dashboard()
		else:
			messagebox.showerror("Error", "Invalid credentials!")

	def go_back(self):
		if self.previous_page_callback:
			self.previous_page_callback()
		else:
			self.root.quit()

	def open_admin_dashboard(self):
		for widget in self.root.winfo_children():
			widget.destroy()
		self.root.title("Admin Dashboard")

		# Background Image
		bg_image = Image.open("admin_bg.jpg")
		bg_image = bg_image.resize((self.root.winfo_screenwidth(), self.root.winfo_screenheight()), Image.LANCZOS)
		self.bg_photo = ImageTk.PhotoImage(bg_image)
		bg_label = tk.Label(self.root, image=self.bg_photo)
		bg_label.place(relwidth=1, relheight=1)

		# Dashboard Frame
		dashboard_frame = tk.Frame(self.root, bg="white", padx=20, pady=20, 
						   bd=5, relief="ridge")  # Add border
		dashboard_frame.place(relx=0.2, rely=0.5, anchor=tk.CENTER)


		# Buttons Frame
		button_frame = tk.Frame(dashboard_frame, bg="white")
		button_frame.pack()

		button_style = {"font": ("Arial", 12, "bold"), "bg": "navyblue", "fg": "white", "width": 25, "height": 1, "bd": 3, "relief": tk.RAISED}
		
		tk.Button(button_frame, text="Approve Customer Requests", command=self.approve_requests, **button_style).pack(pady=10)
		tk.Button(button_frame, text="Approve ATM Requests", command=self.approve_atm_requests, **button_style).pack(pady=10)
		tk.Button(button_frame, text="Generate Transaction Report", command=self.generate_transaction_report, **button_style).pack(pady=10)
		tk.Button(button_frame, text="View Pending Loan Requests", command=self.open_loan_approval, **button_style).pack(pady=10)
		tk.Button(button_frame, text="Logout", command=self.create_login_ui, font=("Arial", 12, "bold"), bg="red", fg="white", width=20, height=1, bd=3, relief=tk.RAISED).pack(pady=10)


	def approve_requests(self):
		connection = create_connection()
		cursor = connection.cursor()
		
		cursor.execute("SELECT id, name, pin, mobile_number, dob, aadhar, nationality FROM account_requests WHERE status = 'Pending'")
		requests = cursor.fetchall()

		if not requests:
			messagebox.showinfo("No Requests", "No pending requests!")
			return

		for request_id, name, pin, mobile_number, dob, aadhar, nationality in requests:
			account_number = random.randint(10000000, 99999999)
			cursor.execute("""
				INSERT INTO customers (account_number, pin, balance, creation_date, name, dob, mobile_number, nationality)
				VALUES (%s, %s, 0.00, CURDATE(), %s, %s, %s, %s)
			""", (account_number, pin, name, dob, mobile_number, nationality))

			cursor.execute("UPDATE account_requests SET account_number = %s, status = 'Approved' WHERE id = %s",
						   (account_number, request_id))

		connection.commit()
		cursor.close()
		connection.close()
		messagebox.showinfo("Success", "All requests approved with new account numbers.")

	# Other methods (approve_atm_requests, generate_transaction_report, open_loan_approval) remain unchanged

	def approve_atm_requests(self):
		connection = create_connection()
		cursor = connection.cursor()
		cursor.execute("SELECT id, account_number FROM atm_requests WHERE status = 'Pending'")
		requests = cursor.fetchall()
	
		if not requests:
			messagebox.showinfo("No Requests", "No pending ATM requests!")
			return

		for request_id, account_number in requests:
			cursor.execute("SELECT name, pin FROM customers WHERE account_number = %s", (account_number,))
			customer_data = cursor.fetchone()
		
			if not customer_data:
				continue 

			name, pin = customer_data
			confirm = messagebox.askyesno("Approve ATM Request", f"Approve ATM request for {name} (Account: {account_number})?")
			if confirm:
				cursor.execute("INSERT INTO atm_cards (account_number, status) VALUES (%s, 'Active')", (account_number,))
				cursor.execute("UPDATE atm_requests SET status = 'Approved' WHERE id = %s", (request_id,))
				self.generate_atm_card(account_number, name, pin)

		connection.commit()
		cursor.close()
		connection.close()
		messagebox.showinfo("Success", "Approved ATM requests. ATM cards generated.")

	def generate_atm_card(self, account_number, name, pin):
		card_template_path = "atm.jpg"  # Path to the ATM card background image
		dpi = 300  # High resolution

		try:
			card = Image.open(card_template_path).convert("RGB")  
			draw = ImageDraw.Draw(card)

			font_path = "arial.ttf"  # Change this if needed
			try:
				font = ImageFont.truetype(font_path, 30)
			except IOError:
				font = ImageFont.load_default()

			# Draw text on the card
			draw.text((50, 20), "ATM CARD", fill="white", font=font)
			draw.text((50, 70), f"Name: {name}", fill="white", font=font)
			draw.text((50, 120), f"Account: {account_number}", fill="white", font=font)
			draw.text((50, 170), f"PIN: {pin}", fill="white", font=font)

			# Save the ATM card with the account number as filename
			save_path = os.path.join(os.getcwd(), f"ATM_Card_{account_number}.png")
			card.save(save_path)
			print(f"ATM Card saved as {save_path}")

		except FileNotFoundError:
			print("Error: ATM card template image 'atm.jpg' not found. Make sure the file exists in the directory.")

	def approve_loan_requests(self):
		connection = create_connection()
		cursor = connection.cursor()

		# Retrieve loan requests with account number
		cursor.execute("""
			SELECT lr.id, lr.account_number, lr.annual_income, lr.amount, c.dob  
			FROM loan_requests lr
			JOIN customers c ON lr.account_number = c.account_number
			WHERE lr.status = 'Pending'
		""")
		requests = cursor.fetchall()

		if not requests:
			messagebox.showinfo("No Requests", "No pending loan requests!")
			return

		for request_id, account_number, annual_income, amount, dob in requests:
			annual_income = float(annual_income)  # Convert Decimal to float
			amount = float(amount)  # Convert Decimal to float
		
			# Calculate age using DOB
			cursor.execute("SELECT TIMESTAMPDIFF(YEAR, %s, CURDATE())", (dob,))
			age = cursor.fetchone()[0]

			max_eligible_amount = 0.4 * annual_income  # Max loan is 40% of annual income

			if age >= 21 and annual_income >= 200000 and amount <= max_eligible_amount:
				cursor.execute("UPDATE loan_requests SET status = 'Approved' WHERE id = %s", (request_id,))
			else:
				cursor.execute("UPDATE loan_requests SET status = 'Rejected' WHERE id = %s", (request_id,))
	
		connection.commit()
		cursor.close()
		connection.close()
		messagebox.showinfo("Success", "Loan requests processed.")

	def generate_transaction_report(self):
		connection = create_connection()
		cursor = connection.cursor()
		cursor.execute("SELECT id, account_number FROM report_requests WHERE status = 'Pending'")
		requests = cursor.fetchall()
	
		if not requests:
			messagebox.showinfo("No Requests", "No pending report requests!")
			return

		for request_id, account_number in requests:
			cursor.execute("SELECT transaction_id, date, Transaction_Type, amount, balance FROM transactions WHERE account_number = %s", (account_number,))
			transactions = cursor.fetchall()
		
			if transactions:
				doc = Document()
				doc.add_heading(f'Transaction Report for Account {account_number}', level=1)

				table = doc.add_table(rows=1, cols=5)
				table.style = 'Table Grid'
				hdr_cells = table.rows[0].cells
				hdr_cells[0].text = 'Transaction ID'
				hdr_cells[1].text = 'Date'
				hdr_cells[2].text = 'Transaction_Type'
				hdr_cells[3].text = 'Amount'
				hdr_cells[4].text = 'Balance'

				for transaction in transactions:
					row_cells = table.add_row().cells
					row_cells[0].text = str(transaction[0])  # Transaction ID
					row_cells[1].text = str(transaction[1])  # Date
					row_cells[2].text = str(transaction[2])  # Type (Deposit/Withdraw)
					row_cells[3].text = str(transaction[3])  # Amount
					row_cells[4].text = str(transaction[4])  # Balance

				# Save the document
				save_path = os.path.join(os.getcwd(), f"Transaction_Report_{account_number}.docx")
				doc.save(save_path)

				# Update report request status in database
				cursor.execute("UPDATE report_requests SET status = 'Approved' WHERE id = %s", (request_id,))

				messagebox.showinfo("Report Generated", f"Transaction report saved: {save_path}")
			else:
				messagebox.showinfo("No Transactions", f"No transactions found for account {account_number}.")
	
		connection.commit()
		cursor.close()
		connection.close()

	def open_loan_approval(self):
		loan_approval_window = tk.Toplevel(self.root)
		LoanApproval(loan_approval_window)

class LoanApproval:
	def __init__(self, window=None):
		self.master = window
		window.geometry("600x400")
		window.title("Loan Approval")

		tk.Label(window, text="Pending Loan Requests", font=("Arial", 14, "bold")).pack(pady=10)

		self.frame = tk.Frame(window)
		self.frame.pack(fill="both", expand=True)

		self.load_loan_requests()

	def load_loan_requests(self):
		connection = create_connection()
		cursor = connection.cursor()

		cursor.execute("""
			SELECT lr.id, lr.account_number, lr.amount, c.annual_income, c.dob  
			FROM loan_requests lr
			JOIN customers c ON lr.account_number = c.account_number
			WHERE lr.status = 'Pending'
		""")
		requests = cursor.fetchall()

		cursor.close()
		connection.close()

		if not requests:
			tk.Label(self.frame, text="No pending loan requests!", font=("Arial", 12)).pack(pady=10)
			return

		for request_id, account_number, amount, annual_income, dob in requests:
			self.create_request_row(request_id, account_number, amount, annual_income, dob)

	def create_request_row(self, request_id, account_number, amount, annual_income, dob):
		frame = tk.Frame(self.frame)
		frame.pack(pady=5, fill="x")

		tk.Label(frame, text=f"Acc No: {account_number}, Loan: ₹{amount}, Income: ₹{annual_income}, DOB: {dob}",
				 font=("Arial", 10), anchor="w").pack(side="left", padx=5)

		# Automatically approve or reject based on criteria
		if self.can_approve_loan(annual_income, amount, dob):
			self.approve_loan(request_id)
			tk.Label(frame, text="Approved", bg="green", fg="white").pack(side="left", padx=5)
		else:
			self.reject_loan(request_id)
			tk.Label(frame, text="Rejected", bg="red", fg="white").pack(side="left", padx=5)

	def can_approve_loan(self, annual_income, amount, dob):
	# Calculate age
		if isinstance(dob, datetime):  # Fix: should be datetime.datetime
			age = (datetime.now() - dob).days // 365  # Calculate age in years
		elif isinstance(dob, str):
			dob = datetime.strptime(dob, "%Y-%m-%d")  # Convert string to datetime
			age = (datetime.now() - dob).days // 365
		elif isinstance(dob, date):  # Fix: Directly reference 'date'
			age = (datetime.now() - datetime.combine(dob, datetime.min.time())).days // 365
		else:
			raise ValueError("DOB must be a datetime, date, or string.")

		max_eligible_amount = 0.4 * annual_income  # Max loan is 40% of annual income

	# Check if the loan can be approved
		if age >= 21 and annual_income >= 200000 and amount <= max_eligible_amount:
			return True
		else:
			messagebox.showinfo("Loan Eligibility", "User is not eligible for loan request.")
			return False


	def approve_loan(self, request_id):
		connection = create_connection()
		cursor = connection.cursor()
		cursor.execute("UPDATE loan_requests SET status = 'Approved' WHERE id = %s", (request_id,))
		connection.commit()
		cursor.close()
		connection.close()

	def reject_loan(self, request_id):
		connection = create_connection()
		cursor = connection.cursor()
		cursor.execute("UPDATE loan_requests SET status = 'Rejected' WHERE id = %s", (request_id,))
		connection.commit()
		cursor.close()
		connection.close()

	def refresh_ui(self):
		for widget in self.frame.winfo_children():
			widget.destroy()
		self.load_loan_requests()
if __name__ == "__main__":
	root = tk.Tk()
	AdminLogin(root)
	root.mainloop()